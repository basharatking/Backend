import os
import io
import uuid
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware

# ── PDF libraries ───────────────────────────────────────────────
import pypdf
import pdfplumber
from PIL import Image
import img2pdf
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import fitz  # PyMuPDF — for compress, rotate, watermark, jpg conversion

# ───────────────────────────────────────────────────────────────
app = FastAPI(title="LivePDF API", version="2.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],          # tighten to your Netlify URL in production
    allow_methods=["*"],
    allow_headers=["*"],
)

MAX_SIZE = 50 * 1024 * 1024  # 50 MB

TMPDIR = Path(tempfile.gettempdir()) / "livepdf"
TMPDIR.mkdir(exist_ok=True)


# ── Helpers ─────────────────────────────────────────────────────

def tmp_path(suffix: str) -> Path:
    return TMPDIR / f"{uuid.uuid4().hex}{suffix}"


def stream_file(path: Path, media_type: str, filename: str) -> StreamingResponse:
    def iterfile():
        with open(path, "rb") as f:
            yield from f
        path.unlink(missing_ok=True)

    return StreamingResponse(
        iterfile(),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def stream_bytes(data: bytes, media_type: str, filename: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def read_upload(upload: UploadFile) -> bytes:
    data = await upload.read()
    if len(data) > MAX_SIZE:
        raise HTTPException(400, "File exceeds 50 MB limit.")
    return data


# ════════════════════════════════════════════════════════════════
# 1.  MERGE PDF
# ════════════════════════════════════════════════════════════════
@app.post("/merge-pdf")
async def merge_pdf(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(400, "Provide at least 2 PDF files.")
    writer = pypdf.PdfWriter()
    for upload in files:
        data = await read_upload(upload)
        reader = pypdf.PdfReader(io.BytesIO(data))
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return stream_bytes(out.getvalue(), "application/pdf", "merged.pdf")


# ════════════════════════════════════════════════════════════════
# 2.  SPLIT PDF
# ════════════════════════════════════════════════════════════════
@app.post("/split-pdf")
async def split_pdf(
    file: UploadFile = File(...),
    mode: str = Form("each"),           # "each" | "range"
    start_page: int = Form(1),
    end_page: int = Form(1),
):
    data = await read_upload(file)
    reader = pypdf.PdfReader(io.BytesIO(data))
    total = len(reader.pages)

    if mode == "range":
        s = max(1, start_page) - 1
        e = min(total, end_page)
        if s >= e:
            raise HTTPException(400, f"Invalid page range. PDF has {total} pages.")
        writer = pypdf.PdfWriter()
        for i in range(s, e):
            writer.add_page(reader.pages[i])
        out = io.BytesIO()
        writer.write(out)
        fname = file.filename.replace(".pdf", "") + f"_pages_{s+1}-{e}.pdf"
        return stream_bytes(out.getvalue(), "application/pdf", fname)

    # mode == "each" → ZIP
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i in range(total):
            writer = pypdf.PdfWriter()
            writer.add_page(reader.pages[i])
            page_buf = io.BytesIO()
            writer.write(page_buf)
            zf.writestr(f"page_{str(i+1).zfill(3)}.pdf", page_buf.getvalue())
    return stream_bytes(zip_buf.getvalue(), "application/zip", "split_pages.zip")


# ════════════════════════════════════════════════════════════════
# 3.  COMPRESS PDF   (PyMuPDF — real image recompression)
# ════════════════════════════════════════════════════════════════
@app.post("/compress-pdf")
async def compress_pdf(
    file: UploadFile = File(...),
    level: str = Form("medium"),   # "low" | "medium" | "high"
):
    data = await read_upload(file)
    doc = fitz.open(stream=data, filetype="pdf")

    # Image quality map
    quality_map = {"low": 85, "medium": 65, "high": 40}
    jpg_quality = quality_map.get(level, 65)

    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                base_img = doc.extract_image(xref)
                img_bytes = base_img["image"]
                pil_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                # Downscale if large
                w, h = pil_img.size
                if w > 1200:
                    pil_img = pil_img.resize((1200, int(h * 1200 / w)), Image.LANCZOS)
                buf = io.BytesIO()
                pil_img.save(buf, format="JPEG", quality=jpg_quality, optimize=True)
                doc.update_stream(xref, buf.getvalue())
            except Exception:
                pass  # skip images that can't be recompressed

    out_buf = io.BytesIO()
    doc.save(out_buf, garbage=4, deflate=True, clean=True)
    doc.close()
    fname = "compressed_" + file.filename
    return stream_bytes(out_buf.getvalue(), "application/pdf", fname)


# ════════════════════════════════════════════════════════════════
# 4.  PDF TO WORD
# ════════════════════════════════════════════════════════════════
@app.post("/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    data = await read_upload(file)
    doc = Document()

    # Heading style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        total = len(pdf.pages)
        if total == 0:
            raise HTTPException(400, "PDF has no pages.")

        for i, page in enumerate(pdf.pages):
            # Page heading
            heading = doc.add_heading(f"Page {i + 1}", level=1)
            heading.runs[0].font.color.rgb = RGBColor(0xD9, 0x3A, 0x1E)

            text = page.extract_text() or ""
            if text.strip():
                for line in text.split("\n"):
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(2)
            else:
                doc.add_paragraph("[No extractable text on this page]").italic = True

            if i < total - 1:
                doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    fname = file.filename.replace(".pdf", "").replace(".PDF", "") + ".docx"
    return stream_bytes(
        out.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        fname,
    )


# ════════════════════════════════════════════════════════════════
# 5.  PDF TO EXCEL   (pdfplumber table detection)
# ════════════════════════════════════════════════════════════════
@app.post("/pdf-to-excel")
async def pdf_to_excel(
    file: UploadFile = File(...),
    mode: str = Form("tables"),   # "tables" | "text"
):
    data = await read_upload(file)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    header_fill = PatternFill("solid", fgColor="D93A1E")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def style_header_row(ws, row_idx: int, ncols: int):
        for col in range(1, ncols + 1):
            cell = ws.cell(row=row_idx, column=col)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_align

    def auto_col_widths(ws):
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    max_len = max(max_len, len(str(cell.value or "")))
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = min(max(max_len + 2, 10), 50)

    sheet_count = 0

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            if mode == "tables":
                tables = page.extract_tables()
                if tables:
                    for t_idx, table in enumerate(tables):
                        if not table:
                            continue
                        sheet_name = f"P{page_num}_T{t_idx+1}"[:31]
                        ws = wb.create_sheet(title=sheet_name)
                        sheet_count += 1
                        for r_idx, row in enumerate(table, start=1):
                            for c_idx, cell_val in enumerate(row, start=1):
                                ws.cell(row=r_idx, column=c_idx, value=cell_val or "")
                            if r_idx == 1:
                                style_header_row(ws, 1, len(row))
                        auto_col_widths(ws)
                else:
                    # Fall back to text extraction for this page
                    text = page.extract_text() or ""
                    if text.strip():
                        sheet_name = f"Page_{page_num}"[:31]
                        ws = wb.create_sheet(title=sheet_name)
                        sheet_count += 1
                        ws.cell(row=1, column=1, value="Line")
                        ws.cell(row=1, column=2, value="Text")
                        style_header_row(ws, 1, 2)
                        for i, line in enumerate(text.split("\n"), start=2):
                            if line.strip():
                                ws.cell(row=i, column=1, value=i - 1)
                                ws.cell(row=i, column=2, value=line.strip())
                        auto_col_widths(ws)

            else:  # text mode — one sheet per page
                text = page.extract_text() or ""
                sheet_name = f"Page_{page_num}"[:31]
                ws = wb.create_sheet(title=sheet_name)
                sheet_count += 1
                ws.cell(row=1, column=1, value="Line")
                ws.cell(row=1, column=2, value="Text")
                style_header_row(ws, 1, 2)
                for i, line in enumerate(text.split("\n"), start=2):
                    if line.strip():
                        ws.cell(row=i, column=1, value=i - 1)
                        ws.cell(row=i, column=2, value=line.strip())
                auto_col_widths(ws)

    if sheet_count == 0:
        ws = wb.create_sheet(title="No Data")
        ws.cell(row=1, column=1, value="No extractable text or tables found in this PDF.")

    out = io.BytesIO()
    wb.save(out)
    fname = file.filename.replace(".pdf", "").replace(".PDF", "") + ".xlsx"
    return stream_bytes(
        out.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        fname,
    )


# ════════════════════════════════════════════════════════════════
# 6.  PDF TO JPG
# ════════════════════════════════════════════════════════════════
@app.post("/pdf-to-jpg")
async def pdf_to_jpg(
    file: UploadFile = File(...),
    dpi: int = Form(150),
):
    data = await read_upload(file)
    doc = fitz.open(stream=data, filetype="pdf")
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("jpeg")
            zf.writestr(f"page_{str(i+1).zfill(3)}.jpg", img_bytes)

    doc.close()
    fname = file.filename.replace(".pdf", "").replace(".PDF", "") + "_images.zip"
    return stream_bytes(zip_buf.getvalue(), "application/zip", fname)


# ════════════════════════════════════════════════════════════════
# 7.  JPG / PNG / TXT → PDF
# ════════════════════════════════════════════════════════════════
@app.post("/jpg-to-pdf")
async def jpg_to_pdf(files: list[UploadFile] = File(...)):
    """Accepts JPG, PNG, TXT files and converts them to a single PDF."""
    images = []
    txt_pages = []

    for upload in files:
        raw = await read_upload(upload)
        ext = Path(upload.filename).suffix.lower()
        if ext in {".jpg", ".jpeg", ".png"}:
            img = Image.open(io.BytesIO(raw)).convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=92)
            images.append(buf.getvalue())
        elif ext == ".txt":
            txt_pages.append(raw.decode("utf-8", errors="replace"))
        else:
            raise HTTPException(400, f"Unsupported file type: {upload.filename}")

    out_buf = io.BytesIO()

    if images and not txt_pages:
        # Pure image conversion — use img2pdf for lossless embedding
        out_buf.write(img2pdf.convert(images))
    else:
        # Mixed or text — use ReportLab
        c = rl_canvas.Canvas(out_buf, pagesize=A4)
        W, H = A4

        for img_bytes in images:
            pil_img = Image.open(io.BytesIO(img_bytes))
            iw, ih = pil_img.size
            ratio = min(W / iw, H / ih, 1.0)
            draw_w, draw_h = iw * ratio, ih * ratio
            x = (W - draw_w) / 2
            y = (H - draw_h) / 2
            c.drawImage(ImageReader(io.BytesIO(img_bytes)), x, y, draw_w, draw_h)
            c.showPage()

        for txt in txt_pages:
            c.setFont("Helvetica", 11)
            margin = 50
            line_h = 16
            y_pos = H - margin
            for line in txt.split("\n"):
                if y_pos < margin + line_h:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y_pos = H - margin
                c.drawString(margin, y_pos, line[:110])
                y_pos -= line_h
            c.showPage()

        c.save()

    fname = "converted.pdf"
    return stream_bytes(out_buf.getvalue(), "application/pdf", fname)


# ════════════════════════════════════════════════════════════════
# 8.  ROTATE PDF
# ════════════════════════════════════════════════════════════════
@app.post("/rotate-pdf")
async def rotate_pdf(
    file: UploadFile = File(...),
    angle: int = Form(90),              # 90 | 180 | 270
    pages: str = Form("all"),           # "all" | "odd" | "even"
):
    data = await read_upload(file)
    doc = fitz.open(stream=data, filetype="pdf")
    total = len(doc)

    for i, page in enumerate(doc):
        page_num = i + 1
        if pages == "odd" and page_num % 2 == 0:
            continue
        if pages == "even" and page_num % 2 != 0:
            continue
        page.set_rotation((page.rotation + angle) % 360)

    out = io.BytesIO()
    doc.save(out)
    doc.close()
    fname = "rotated_" + file.filename
    return stream_bytes(out.getvalue(), "application/pdf", fname)


# ════════════════════════════════════════════════════════════════
# 9.  UNLOCK PDF  (remove password / encryption)
# ════════════════════════════════════════════════════════════════
@app.post("/unlock-pdf")
async def unlock_pdf(
    file: UploadFile = File(...),
    password: str = Form(""),
):
    data = await read_upload(file)
    doc = fitz.open(stream=data, filetype="pdf")

    if doc.is_encrypted:
        ok = doc.authenticate(password)
        if not ok:
            raise HTTPException(400, "Wrong password — could not unlock PDF.")

    out = io.BytesIO()
    doc.save(out, encryption=fitz.PDF_ENCRYPT_NONE)
    doc.close()
    fname = "unlocked_" + file.filename
    return stream_bytes(out.getvalue(), "application/pdf", fname)


# ════════════════════════════════════════════════════════════════
# 10. ADD WATERMARK
# ════════════════════════════════════════════════════════════════
@app.post("/add-watermark")
async def add_watermark(
    file: UploadFile = File(...),
    text: str = Form("CONFIDENTIAL"),
    opacity: float = Form(0.2),
    position: str = Form("center"),  # "center" | "top" | "bottom"
):
    data = await read_upload(file)
    doc = fitz.open(stream=data, filetype="pdf")

    alpha_int = max(1, min(255, int(opacity * 255)))

    for page in doc:
        w, h = page.rect.width, page.rect.height
        font_size = min(w, h) * 0.08

        if position == "center":
            # Diagonal watermark
            page.insert_text(
                fitz.Point(w * 0.15, h * 0.55),
                text,
                fontsize=font_size,
                rotate=45,
                color=(0.5, 0.5, 0.5),
                fill_opacity=opacity,
                overlay=True,
            )
        elif position == "top":
            page.insert_text(
                fitz.Point(w * 0.5 - len(text) * font_size * 0.25, h - font_size - 20),
                text,
                fontsize=font_size,
                color=(0.5, 0.5, 0.5),
                fill_opacity=opacity,
                overlay=True,
            )
        else:  # bottom
            page.insert_text(
                fitz.Point(w * 0.5 - len(text) * font_size * 0.25, font_size + 20),
                text,
                fontsize=font_size,
                color=(0.5, 0.5, 0.5),
                fill_opacity=opacity,
                overlay=True,
            )

    out = io.BytesIO()
    doc.save(out)
    doc.close()
    fname = "watermarked_" + file.filename
    return stream_bytes(out.getvalue(), "application/pdf", fname)


# ════════════════════════════════════════════════════════════════
# Health check
# ════════════════════════════════════════════════════════════════
@app.get("/health")
def health():
    return {"status": "ok", "version": "2.0"}
